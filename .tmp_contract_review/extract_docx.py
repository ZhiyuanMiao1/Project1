from docx import Document
import sys

for path in sys.argv[1:]:
    print(f"\n===== {path} =====")
    doc = Document(path)
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if text:
            print(f"P{i}: {text}")
    for ti, table in enumerate(doc.tables, 1):
        print(f"--- TABLE {ti} ---")
        for ri, row in enumerate(table.rows, 1):
            cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
            print(f"R{ri}: " + " || ".join(cells))
